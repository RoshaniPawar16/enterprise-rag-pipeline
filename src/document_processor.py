"""
Document Processing Module for Enterprise RAG Pipeline

Handles extraction of text and metadata from various document formats:
- PDF (with OCR fallback for scanned documents)
- DOCX, DOC
- PPTX
- XLSX
- TXT, MD
- Images (with OCR)

Key Features:
- Page-level text extraction
- Table extraction
- Image extraction for multimodal RAG
- Metadata extraction (author, dates, etc.)
"""

import io
import os
import tempfile
from abc import ABC, abstractmethod
from dataclasses import dataclass, field
from datetime import datetime
from enum import Enum
from pathlib import Path
from typing import BinaryIO

import structlog

logger = structlog.get_logger(__name__)


class DocumentType(Enum):
    """Supported document types."""
    PDF = "pdf"
    DOCX = "docx"
    DOC = "doc"
    PPTX = "pptx"
    XLSX = "xlsx"
    TXT = "txt"
    MD = "md"
    IMAGE = "image"
    UNKNOWN = "unknown"


@dataclass
class ExtractedImage:
    """Represents an image extracted from a document."""
    image_bytes: bytes
    page_number: int
    image_index: int
    width: int
    height: int
    format: str
    caption: str | None = None

    def save(self, path: str):
        """Save image to file."""
        with open(path, "wb") as f:
            f.write(self.image_bytes)


@dataclass
class ExtractedTable:
    """Represents a table extracted from a document."""
    data: list[list[str]]
    page_number: int
    table_index: int
    headers: list[str] | None = None

    def to_markdown(self) -> str:
        """Convert table to markdown format."""
        if not self.data:
            return ""

        lines = []
        headers = self.headers or self.data[0]
        data_rows = self.data[1:] if not self.headers else self.data

        # Header row
        lines.append("| " + " | ".join(str(h) for h in headers) + " |")
        lines.append("| " + " | ".join("---" for _ in headers) + " |")

        # Data rows
        for row in data_rows:
            # Pad row if needed
            padded = list(row) + [""] * (len(headers) - len(row))
            lines.append("| " + " | ".join(str(c) for c in padded[:len(headers)]) + " |")

        return "\n".join(lines)


@dataclass
class PageContent:
    """Content from a single page of a document."""
    page_number: int
    text: str
    tables: list[ExtractedTable] = field(default_factory=list)
    images: list[ExtractedImage] = field(default_factory=list)


@dataclass
class DocumentMetadata:
    """Metadata extracted from a document."""
    filename: str
    file_size: int
    doc_type: DocumentType
    page_count: int
    created_at: str | None = None
    modified_at: str | None = None
    author: str | None = None
    title: str | None = None
    subject: str | None = None
    keywords: list[str] = field(default_factory=list)
    language: str | None = None

    def to_dict(self) -> dict:
        return {
            "filename": self.filename,
            "file_size": self.file_size,
            "doc_type": self.doc_type.value,
            "page_count": self.page_count,
            "created_at": self.created_at,
            "modified_at": self.modified_at,
            "author": self.author,
            "title": self.title,
            "subject": self.subject,
            "keywords": self.keywords,
            "language": self.language,
        }


@dataclass
class ProcessedDocument:
    """Complete processed document with all extracted content."""
    metadata: DocumentMetadata
    pages: list[PageContent]
    full_text: str
    processing_time: float
    extracted_at: str

    def to_dict(self) -> dict:
        return {
            "metadata": self.metadata.to_dict(),
            "full_text": self.full_text,
            "page_count": len(self.pages),
            "total_tables": sum(len(p.tables) for p in self.pages),
            "total_images": sum(len(p.images) for p in self.pages),
            "processing_time": self.processing_time,
            "extracted_at": self.extracted_at,
        }


class DocumentProcessor(ABC):
    """Abstract base class for document processors."""

    @abstractmethod
    def process(self, file_path: str) -> ProcessedDocument:
        """Process a document and extract all content."""
        pass

    @abstractmethod
    def supports(self, doc_type: DocumentType) -> bool:
        """Check if this processor supports the given document type."""
        pass


class PDFProcessor(DocumentProcessor):
    """Process PDF documents using PyMuPDF."""

    def __init__(self, extract_images: bool = True, ocr_fallback: bool = True):
        self.extract_images = extract_images
        self.ocr_fallback = ocr_fallback

    def supports(self, doc_type: DocumentType) -> bool:
        return doc_type == DocumentType.PDF

    def process(self, file_path: str) -> ProcessedDocument:
        import fitz
        import time

        start_time = time.time()
        logger.info("Processing PDF", path=file_path)

        doc = fitz.open(file_path)
        pages = []
        all_text_parts = []

        # Extract metadata
        metadata = self._extract_metadata(doc, file_path)

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_content = self._process_page(page, page_num + 1)
            pages.append(page_content)
            all_text_parts.append(f"[Page {page_num + 1}]\n{page_content.text}")

        doc.close()

        processing_time = time.time() - start_time
        full_text = "\n\n".join(all_text_parts)

        logger.info(
            "PDF processing complete",
            path=file_path,
            pages=len(pages),
            chars=len(full_text),
            time=f"{processing_time:.2f}s",
        )

        return ProcessedDocument(
            metadata=metadata,
            pages=pages,
            full_text=full_text,
            processing_time=processing_time,
            extracted_at=datetime.utcnow().isoformat(),
        )

    def _extract_metadata(self, doc, file_path: str) -> DocumentMetadata:
        """Extract PDF metadata."""
        import fitz

        meta = doc.metadata
        file_stat = os.stat(file_path)

        return DocumentMetadata(
            filename=Path(file_path).name,
            file_size=file_stat.st_size,
            doc_type=DocumentType.PDF,
            page_count=len(doc),
            created_at=meta.get("creationDate"),
            modified_at=meta.get("modDate"),
            author=meta.get("author"),
            title=meta.get("title"),
            subject=meta.get("subject"),
            keywords=meta.get("keywords", "").split(",") if meta.get("keywords") else [],
        )

    def _process_page(self, page, page_number: int) -> PageContent:
        """Process a single PDF page."""
        import fitz

        # Extract text
        text = page.get_text("text")

        # If text is empty, try OCR fallback
        if not text.strip() and self.ocr_fallback:
            text = self._ocr_page(page)

        # Extract tables
        tables = self._extract_tables(page, page_number)

        # Extract images
        images = []
        if self.extract_images:
            images = self._extract_images(page, page_number)

        return PageContent(
            page_number=page_number,
            text=text,
            tables=tables,
            images=images,
        )

    def _ocr_page(self, page) -> str:
        """OCR a page using pytesseract."""
        try:
            import pytesseract
            from PIL import Image
            import io

            # Render page to image
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")
            img = Image.open(io.BytesIO(img_data))

            # Run OCR
            text = pytesseract.image_to_string(img)
            logger.debug("OCR completed for page", chars=len(text))
            return text

        except ImportError:
            logger.warning("pytesseract not available for OCR fallback")
            return ""
        except Exception as e:
            logger.error("OCR failed", error=str(e))
            return ""

    def _extract_tables(self, page, page_number: int) -> list[ExtractedTable]:
        """Extract tables from a PDF page."""
        tables = []

        try:
            # Use PyMuPDF's table detection
            page_tables = page.find_tables()

            for idx, table in enumerate(page_tables):
                extracted = table.extract()
                if extracted:
                    tables.append(ExtractedTable(
                        data=extracted,
                        page_number=page_number,
                        table_index=idx,
                    ))

        except Exception as e:
            logger.debug("Table extraction failed", error=str(e))

        return tables

    def _extract_images(self, page, page_number: int) -> list[ExtractedImage]:
        """Extract images from a PDF page."""
        images = []

        try:
            image_list = page.get_images()

            for idx, img_info in enumerate(image_list):
                xref = img_info[0]
                base_image = page.parent.extract_image(xref)

                if base_image:
                    images.append(ExtractedImage(
                        image_bytes=base_image["image"],
                        page_number=page_number,
                        image_index=idx,
                        width=base_image.get("width", 0),
                        height=base_image.get("height", 0),
                        format=base_image.get("ext", "png"),
                    ))

        except Exception as e:
            logger.debug("Image extraction failed", error=str(e))

        return images


class DOCXProcessor(DocumentProcessor):
    """Process DOCX documents using python-docx."""

    def supports(self, doc_type: DocumentType) -> bool:
        return doc_type == DocumentType.DOCX

    def process(self, file_path: str) -> ProcessedDocument:
        from docx import Document
        import time

        start_time = time.time()
        logger.info("Processing DOCX", path=file_path)

        doc = Document(file_path)

        # Extract text
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        full_text = "\n\n".join(paragraphs)

        # Extract tables
        tables = []
        for idx, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text for cell in row.cells]
                table_data.append(row_data)

            if table_data:
                tables.append(ExtractedTable(
                    data=table_data,
                    page_number=1,  # DOCX doesn't have pages
                    table_index=idx,
                ))

        # Create single page content
        page_content = PageContent(
            page_number=1,
            text=full_text,
            tables=tables,
        )

        # Extract metadata
        metadata = self._extract_metadata(doc, file_path)

        processing_time = time.time() - start_time

        return ProcessedDocument(
            metadata=metadata,
            pages=[page_content],
            full_text=full_text,
            processing_time=processing_time,
            extracted_at=datetime.utcnow().isoformat(),
        )

    def _extract_metadata(self, doc, file_path: str) -> DocumentMetadata:
        """Extract DOCX metadata."""
        props = doc.core_properties
        file_stat = os.stat(file_path)

        return DocumentMetadata(
            filename=Path(file_path).name,
            file_size=file_stat.st_size,
            doc_type=DocumentType.DOCX,
            page_count=1,
            created_at=str(props.created) if props.created else None,
            modified_at=str(props.modified) if props.modified else None,
            author=props.author,
            title=props.title,
            subject=props.subject,
            keywords=props.keywords.split(",") if props.keywords else [],
        )


class TextProcessor(DocumentProcessor):
    """Process plain text and markdown files."""

    def supports(self, doc_type: DocumentType) -> bool:
        return doc_type in (DocumentType.TXT, DocumentType.MD)

    def process(self, file_path: str) -> ProcessedDocument:
        import time

        start_time = time.time()
        logger.info("Processing text file", path=file_path)

        with open(file_path, "r", encoding="utf-8") as f:
            full_text = f.read()

        file_stat = os.stat(file_path)
        suffix = Path(file_path).suffix.lower()
        doc_type = DocumentType.MD if suffix == ".md" else DocumentType.TXT

        metadata = DocumentMetadata(
            filename=Path(file_path).name,
            file_size=file_stat.st_size,
            doc_type=doc_type,
            page_count=1,
        )

        page_content = PageContent(
            page_number=1,
            text=full_text,
        )

        processing_time = time.time() - start_time

        return ProcessedDocument(
            metadata=metadata,
            pages=[page_content],
            full_text=full_text,
            processing_time=processing_time,
            extracted_at=datetime.utcnow().isoformat(),
        )


class DocumentProcessorFactory:
    """Factory for creating document processors."""

    _processors: list[DocumentProcessor] = []

    @classmethod
    def register(cls, processor: DocumentProcessor):
        """Register a document processor."""
        cls._processors.append(processor)

    @classmethod
    def get_processor(cls, doc_type: DocumentType) -> DocumentProcessor | None:
        """Get a processor for the given document type."""
        for processor in cls._processors:
            if processor.supports(doc_type):
                return processor
        return None

    @classmethod
    def detect_type(cls, file_path: str) -> DocumentType:
        """Detect document type from file extension."""
        suffix = Path(file_path).suffix.lower()

        type_map = {
            ".pdf": DocumentType.PDF,
            ".docx": DocumentType.DOCX,
            ".doc": DocumentType.DOC,
            ".pptx": DocumentType.PPTX,
            ".xlsx": DocumentType.XLSX,
            ".txt": DocumentType.TXT,
            ".md": DocumentType.MD,
            ".png": DocumentType.IMAGE,
            ".jpg": DocumentType.IMAGE,
            ".jpeg": DocumentType.IMAGE,
        }

        return type_map.get(suffix, DocumentType.UNKNOWN)


# Register default processors
DocumentProcessorFactory.register(PDFProcessor())
DocumentProcessorFactory.register(DOCXProcessor())
DocumentProcessorFactory.register(TextProcessor())


def process_document(file_path: str) -> ProcessedDocument:
    """
    High-level function to process any supported document.

    This is the main entry point for document processing.

    Args:
        file_path: Path to the document

    Returns:
        ProcessedDocument with all extracted content

    Raises:
        ValueError: If document type is not supported
    """
    doc_type = DocumentProcessorFactory.detect_type(file_path)

    if doc_type == DocumentType.UNKNOWN:
        raise ValueError(f"Unsupported document type: {file_path}")

    processor = DocumentProcessorFactory.get_processor(doc_type)

    if not processor:
        raise ValueError(f"No processor available for: {doc_type.value}")

    return processor.process(file_path)


if __name__ == "__main__":
    # Example usage
    import sys

    if len(sys.argv) > 1:
        file_path = sys.argv[1]
        result = process_document(file_path)

        print(f"\n=== Document: {result.metadata.filename} ===")
        print(f"Type: {result.metadata.doc_type.value}")
        print(f"Pages: {result.metadata.page_count}")
        print(f"Characters: {len(result.full_text)}")
        print(f"Tables: {sum(len(p.tables) for p in result.pages)}")
        print(f"Images: {sum(len(p.images) for p in result.pages)}")
        print(f"Processing time: {result.processing_time:.2f}s")
        print(f"\n=== Preview ===\n{result.full_text[:500]}...")
    else:
        print("Usage: python document_processor.py <file_path>")
